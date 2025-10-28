from docx import Document
import re
import os


# Set your root directory containing DOCX files
root_dir = 'Dokumenti'


# Folder to save individual section files
output_dir = 'Extracted_Sections'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)


# Pattern to detect lines starting with a number (e.g., "1.", "2.", etc.)
section_header_pattern = re.compile(r'^\s*\d+\.\s*')


section_global_count = 0  # To ensure unique section numbering across all files


# Function to process each DOCX file
def process_docx(file_path):
    global section_global_count
    print(f"Processing file: {file_path}")
    doc = Document(file_path)


    # Remove empty paragraphs
    non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip() != '']


    # Create a cleaned document with non-empty paragraphs
    clean_doc = Document()
    for p in non_empty_paragraphs:
        new_para = clean_doc.add_paragraph()
        for run in p.runs:
            new_run = new_para.add_run(run.text)
            # Copy formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb
            new_run.style = run.style


    section_count = 0
    current_doc = None


    for paragraph in clean_doc.paragraphs:
        text = paragraph.text.strip()


        # Check if this paragraph is a section header
        if section_header_pattern.match(text):
            # Save previous section if exists
            if current_doc:
                filename = os.path.join(output_dir, f'section_{section_global_count}.docx')
                current_doc.save(filename)
                section_global_count += 1
            # Start a new section
            current_doc = Document()
            section_count += 1
            # Add the header paragraph
            new_paragraph = current_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                # Copy formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb
                new_run.style = run.style
        else:
            # Continue adding paragraph to current section
            if current_doc:
                new_paragraph = current_doc.add_paragraph()
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    # Copy formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.color.rgb = run.font.color.rgb
                    new_run.style = run.style


    # Save the last section if exists
    if current_doc:
        filename = os.path.join(output_dir, f'section_{section_global_count}.docx')
        current_doc.save(filename)
        section_global_count += 1


# Traverse through the directory and process each .docx file
for dirpath, dirnames, filenames in os.walk(root_dir):
    for filename in filenames:
        if filename.lower().endswith('.docx'):
            file_path = os.path.join(dirpath, filename)
            process_docx(file_path)


print(f"Finished processing all files. Total sections created: {section_global_count}.")